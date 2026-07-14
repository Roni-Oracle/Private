#!/usr/bin/env python3
"""Audit a Hebrew LI-CBT clinical-case draft for length, coverage, and privacy signals."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path


WORD_RE = re.compile(r"[\w\u0590-\u05FF]+(?:[-־'][\w\u0590-\u05FF]+)*", re.UNICODE)

COVERAGE = {
    "setting_consent_privacy": [r"הסכמ", r"סודיות|אנונימ", r"סטאז|סטודנט|הדרכ"],
    "presenting_problem_4w_fff": [r"תלונ|בעיה מרכזית", r"4W|מה.*היכן.*מתי.*מי", r"FFF|Fight|Flight|Freeze|לחימה|בריחה|קיפאון"],
    "risk_scope": [r"אובד|מסוכנות|פגיעה עצמית|סיכון", r"התאמה.*LI|LICBT|גבולות.*הכשר|הפניה"],
    "measurement": [r"PHQ.?9", r"GAD.?7", r"מדד|שאלון|גרף|מגמ"],
    "formulation_goal_plan": [r"המשג|מעגל.*שימור", r"מטרת?.*טיפול|מטרה ממוקדת", r"תוכנית טיפול|תכנית טיפול"],
    "cognitive_work": [r"יומן 1", r"יומן 2", r"יומן 3", r"סוקראט|חץ אנכי|אמונת יסוד"],
    "behavioral_work": [r"חשיפ|הפעלה התנהגותית|ניסוי התנהגותי|פתרון בעיות|אסרטיב"],
    "empathy": [r"אמפת", r"הסכמה|פריקה מנשק", r"שיקוף", r"I Feel|איי פיל", r"ליטוף|עידוד", r"הזמנה לשינוי|שאלה"],
    "outcome_relapse": [r"תוצאות|סיום טיפול|התחלה.*סיום", r"מניעת נפילה|מניעת הישנות|סימני אזהרה"],
    "supervision_reflection": [r"הדרכה|מנחה", r"רפלקציה|בדיעבד|הייתי משנה|למדתי"],
}

PLACEHOLDERS = [r"\[מידע חסר\]", r"\[להשלמה\]", r"TODO", r"XXX", r"___+"]
PRIVACY_PATTERNS = {
    "email": r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b",
    "phone": r"(?<!\d)(?:\+?972[- ]?|0)(?:[23489]|5\d)[- ]?\d{3}[- ]?\d{4}(?!\d)",
    "israeli_id_like": r"(?<!\d)\d{9}(?!\d)",
    "exact_date": r"(?<!\d)(?:0?[1-9]|[12]\d|3[01])[./-](?:0?[1-9]|1[0-2])[./-](?:19|20)\d{2}(?!\d)",
}


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise SystemExit("python-docx is required to audit .docx files") from exc
    doc = Document(path)
    blocks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def read_input(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx(path)
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8-sig")
    raise SystemExit(f"Unsupported input type: {suffix}. Use .md, .txt, or .docx")


def check_patterns(text: str, groups: list[str]) -> dict:
    hits = []
    for pattern in groups:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        hits.append(bool(match))
    return {"passed": all(hits), "matched": sum(hits), "required": len(hits)}


def audit(path: Path) -> dict:
    text = read_input(path)
    words = WORD_RE.findall(text)
    word_count = len(words)
    # Conservative range: headings, tables and Hebrew typography vary greatly.
    estimated_pages = {"low": round(word_count / 520, 1), "high": round(word_count / 380, 1)}
    coverage = {name: check_patterns(text, pats) for name, pats in COVERAGE.items()}
    missing = [name for name, result in coverage.items() if not result["passed"]]
    placeholders = [p for p in PLACEHOLDERS if re.search(p, text, re.IGNORECASE)]
    privacy = {
        name: len(re.findall(pattern, text, re.IGNORECASE))
        for name, pattern in PRIVACY_PATTERNS.items()
    }
    quote_count = len(re.findall(r"[\"“”׳״].{4,120}?[\"“”׳״]", text))
    length_status = "estimated_within_5_8"
    if estimated_pages["high"] < 5:
        length_status = "probably_under_5"
    elif estimated_pages["low"] > 8:
        length_status = "probably_over_8"
    elif word_count < 2200:
        length_status = "likely_thin_or_under_5"
    elif word_count > 4000:
        length_status = "likely_over_8_or_dense"

    critical = {
        "length_roughly_compatible_with_5_8_pages": 2200 <= word_count <= 4000,
        "risk_scope": coverage["risk_scope"]["passed"],
        "measurement": coverage["measurement"]["passed"],
        "formulation_goal_plan": coverage["formulation_goal_plan"]["passed"],
        "privacy_clear_of_obvious_patterns": not any(privacy.values()),
        "no_placeholders": not placeholders,
    }
    return {
        "file": str(path),
        "word_count": word_count,
        "estimated_pages": estimated_pages,
        "length_status": length_status,
        "coverage": coverage,
        "missing_coverage_groups": missing,
        "short_quote_count": quote_count,
        "placeholder_patterns_found": placeholders,
        "privacy_signal_counts": privacy,
        "critical_gates": critical,
        "passes_automated_gate": all(critical.values()) and not missing,
        "note": "Estimate only. Render the final DOCX to verify actual 5–8 pages and inspect false positives manually.",
    }


def print_text(report: dict) -> None:
    print(f"File: {report['file']}")
    print(f"Words: {report['word_count']}")
    ep = report["estimated_pages"]
    print(f"Estimated pages: {ep['low']}–{ep['high']} ({report['length_status']})")
    print("Coverage:")
    for name, result in report["coverage"].items():
        mark = "OK" if result["passed"] else "MISSING/PARTIAL"
        print(f"  {mark:15} {name} ({result['matched']}/{result['required']})")
    print(f"Short quotation-like excerpts: {report['short_quote_count']}")
    print(f"Placeholders: {report['placeholder_patterns_found'] or 'none'}")
    print(f"Privacy signals: {report['privacy_signal_counts']}")
    print(f"Automated gate: {'PASS' if report['passes_automated_gate'] else 'REVIEW REQUIRED'}")
    print(report["note"])


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Clinical-case .md, .txt, or .docx file")
    parser.add_argument("--json", action="store_true", help="Print JSON instead of text")
    args = parser.parse_args()
    if not args.input.is_file():
        parser.error(f"File not found: {args.input}")
    report = audit(args.input)
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print_text(report)
    return 0 if report["passes_automated_gate"] else 1


if __name__ == "__main__":
    sys.exit(main())
